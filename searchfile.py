# -*- coding: utf-8 -*-
"""
Created on Sun Mar  3 09:46:50 2019

@author: yang
"""

from docx import Document
import os,sys

class SearchFile:
    def __init__(self):
        self.word = ""
        self.filename = ""
        self.root_dir = r"../TCM-Retrieval-System/books/"
        self.root = os.path.abspath(self.root_dir)
        
    def get_content(self,word,filename=""):
        self.word = word
        self.filename = filename
        content_list=[]
        filename = filename+".docx"
        try:
            content_list= self.find_files(filename)
        except Exception as e:
            print(e)
        return content_list

        
    def search_word(self,filename):
        #打开文档
        document = Document(filename)
        # document = Document(r'C:\Users\Cheng\Desktop\kword\words\wind.docx')
        #print(filename)
        #读取每段资料
        #l = [ paragraph.text.encode('gb2312') for paragraph in document.paragraphs]
        l = [ paragraph.text.encode('utf-8') for paragraph in document.paragraphs]
        #输出并观察结果，也可以通过其他手段处理文本即可
        wordB = self.word.encode('utf-8')
        wlist = []
        for i in l:
            i=i.strip()
            if i.find(wordB)!=-1 and i!= "":
                wlist.append(i.decode())
                print(len(wlist))
        return wlist
        
    def get_process_files(self,filename = ""):
        """process all files in directory"""
        cur_dir=os.path.abspath(self.root)
        file_list=os.listdir(cur_dir)
        #print(file_list)
        process_list=[]
        if filename != "":
            index = file_list.index(filename)
            targetfile = cur_dir+"\\"+ file_list[index]
            if os.path.isfile(targetfile):
                process_list.append(targetfile)
        else:
            for file in file_list:
                fullfile=cur_dir+"\\"+file
                if os.path.isfile(fullfile):
                    process_list.append(fullfile)
                elif os.path.isdir(fullfile):
                    dir_extra_list=get_process_files(fullfile)
                    if len(dir_extra_list)!=0:
                        for x in dir_extra_list:
                            process_list.append(x)
        return process_list

    def find_files(self,filename=""):
        process_list=self.get_process_files(filename)
        content_list=[]
        for files in process_list:
            #try:
            list0 = self.search_word(files)
            #print(list0)
            content_list = content_list + list0
            #except Exception as e:
             #   print(e)
        return content_list

if __name__=='__main__':
    #文件根目录
    #root_dir=sys.argv[1]
    #root_dir = r"../TCM-Retrieval-System/books/"
    #要搜索的关键字
    #word=sys.argv[2]
    #root = os.path.abspath(root_dir)
    #print(root)
    word = "枸杞"
    filename = "本草纲目"
    sf = SearchFile()
    try:
        list = sf.get_content(word,filename):
            
    except Exception as e:
        print(e)
